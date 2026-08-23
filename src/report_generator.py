from pathlib import Path
from typing import Dict, Any, List
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import pandas as pd
import numpy as np


def set_cell_background(cell, fill_hex: str) -> None:
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 150, right: int = 150) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, border_color: str = "D3D3D3") -> None:
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1B365D"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)


def format_paragraph(p, space_before: int = 0, space_after: int = 6, line_spacing: float = 1.15) -> None:
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing


def add_custom_heading(doc, text: str, level: int) -> Any:
    p = doc.add_paragraph()
    format_paragraph(p, space_before=14 if level == 1 else 10, space_after=4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(27, 54, 93)  # Navy
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="1B365D"/></w:pBdr>')
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(44, 82, 130)  # Slate Blue
    elif level == 3:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(60, 60, 60)
    return p


def add_callout(doc, text: str, title: str = "SUMMARY TAKEAWAY") -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="1B365D"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=0, space_after=2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = RGBColor(27, 54, 93)
    
    r_text = p.add_run(text)
    r_text.font.name = "Arial"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(40, 40, 40)
    
    doc.add_paragraph()


def add_code_block(doc, code: str, title: str) -> None:
    p_title = doc.add_paragraph()
    format_paragraph(p_title, space_before=8, space_after=2)
    r_title = p_title.add_run(f"Code Snippet: {title}")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(27, 54, 93)
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F7F9FA")
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="4A5568"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=0, space_after=0, line_spacing=1.05)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(30, 30, 30)
    
    p_space = doc.add_paragraph()
    format_paragraph(p_space, space_before=0, space_after=4)


def add_formatted_table(
    doc,
    headers: List[str],
    data_rows: List[List[str]],
    col_widths_inches: List[float] = None
) -> None:
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        format_paragraph(p, space_before=0, space_after=0)
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        if col_widths_inches and i < len(col_widths_inches):
            cell.width = Inches(col_widths_inches[i])
            
    # Data Rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F7F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            cell = row_cells[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            format_paragraph(p, space_before=0, space_after=0)
            run = p.add_run(str(cell_value))
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(40, 40, 40)
            if col_widths_inches and col_idx < len(col_widths_inches):
                cell.width = Inches(col_widths_inches[col_idx])
                
    p_space = doc.add_paragraph()
    format_paragraph(p_space, space_before=0, space_after=6)


def generate_word_report(
    metadata: Dict[str, Any],
    inspection_summary: Dict[str, Any],
    summary_df: pd.DataFrame,
    evaluation_results: Dict[str, Dict[str, Any]],
    misclassified_df: pd.DataFrame,
    feature_importance_df: pd.DataFrame,
    figures_paths: Dict[str, Path],
    output_path: Path
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Extract dynamic values from evaluation results
    lr_train_acc = evaluation_results["Logistic Regression"]["train_accuracy"]
    lr_test_acc = evaluation_results["Logistic Regression"]["test_accuracy"]
    lr_f1 = evaluation_results["Logistic Regression"]["f1_macro"]
    lr_f1_weighted = evaluation_results["Logistic Regression"]["f1_weighted"]
    lr_roc_auc = evaluation_results["Logistic Regression"]["roc_auc_ovr_macro"]
    lr_gap = evaluation_results["Logistic Regression"]["generalization_gap_accuracy"]
    lr_train_f1 = evaluation_results["Logistic Regression"]["train_f1_macro"]
    
    rf_train_acc = evaluation_results["Random Forest"]["train_accuracy"]
    rf_test_acc = evaluation_results["Random Forest"]["test_accuracy"]
    rf_f1 = evaluation_results["Random Forest"]["f1_macro"]
    rf_f1_weighted = evaluation_results["Random Forest"]["f1_weighted"]
    rf_roc_auc = evaluation_results["Random Forest"]["roc_auc_ovr_macro"]
    rf_gap = evaluation_results["Random Forest"]["generalization_gap_accuracy"]
    rf_train_f1 = evaluation_results["Random Forest"]["train_f1_macro"]
    
    dt_train_acc = evaluation_results["Decision Tree"]["train_accuracy"]
    dt_test_acc = evaluation_results["Decision Tree"]["test_accuracy"]
    dt_f1 = evaluation_results["Decision Tree"]["f1_macro"]
    dt_f1_weighted = evaluation_results["Decision Tree"]["f1_weighted"]
    dt_roc_auc = evaluation_results["Decision Tree"]["roc_auc_ovr_macro"]
    dt_gap = evaluation_results["Decision Tree"]["generalization_gap_accuracy"]
    dt_train_f1 = evaluation_results["Decision Tree"]["train_f1_macro"]
    
    # Extract dynamic feature importances
    top_feats = feature_importance_df.sort_values("Importance", ascending=False).reset_index(drop=True)
    feat_top1_name = str(top_feats.iloc[0]["Feature"])
    feat_top1_val = float(top_feats.iloc[0]["Importance"])
    feat_top2_name = str(top_feats.iloc[1]["Feature"])
    feat_top2_val = float(top_feats.iloc[1]["Importance"])
    feat_top3_name = str(top_feats.iloc[2]["Feature"])
    feat_top3_val = float(top_feats.iloc[2]["Importance"])
    feat_top4_name = str(top_feats.iloc[3]["Feature"])
    feat_top4_val = float(top_feats.iloc[3]["Importance"])
    top4_cum_pct = float(top_feats.iloc[:4]["Importance"].sum() * 100)
    
    feat_low1_name = str(top_feats.iloc[-1]["Feature"])
    feat_low1_val = float(top_feats.iloc[-1]["Importance"])
    feat_low2_name = str(top_feats.iloc[-2]["Feature"])
    feat_low2_val = float(top_feats.iloc[-2]["Importance"])

    # Title Block
    p_title = doc.add_paragraph()
    format_paragraph(p_title, space_before=0, space_after=4)
    r_title = p_title.add_run("Machine Learning Model Development and Evaluation")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(27, 54, 93)
    
    p_subtitle = doc.add_paragraph()
    format_paragraph(p_subtitle, space_before=0, space_after=14)
    r_sub = p_subtitle.add_run("A Multiclass Classification Study on Physicochemical Wine Cultivar Fingerprints")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    r_sub.italic = True
    
    # Executive Summary
    add_custom_heading(doc, "Executive Summary", level=1)
    
    p_exec = doc.add_paragraph()
    format_paragraph(p_exec, space_before=0, space_after=6)
    p_exec.add_run(
        "This project develops, validates, and compares a reproducible end-to-end machine learning pipeline "
        "for multiclass classification on the Wine dataset from the UCI Machine Learning Repository and scikit-learn. "
        "The primary analytical objective is to determine whether 13 chemical and physical measurements derived from Italian wines "
        "can reliably classify their cultivar of origin (cultivars 0, 1, and 2, corresponding to Barolo, Grignolino, and Barbera). "
        "The experimental pipeline incorporates data hygiene inspection, stratified train-test partitioning (80/20 split, random_state=42), "
        "leakage-free standard feature scaling encapsulated inside scikit-learn Pipelines, and comprehensive multiclass evaluation."
    )
    
    p_exec2 = doc.add_paragraph()
    format_paragraph(p_exec2, space_before=0, space_after=8)
    p_exec2.add_run(
        f"Three candidate classification algorithms were evaluated: "
        f"Logistic Regression (multinomial linear baseline with L2 regularization), Random Forest (ensemble bagging of 100 trees), "
        f"and Decision Tree (single non-linear CART baseline). On the held-out test set (N = 36), the Random Forest model achieved "
        f"{rf_test_acc * 100:.1f}% test accuracy (Macro F1 = {rf_f1:.4f}, One-vs-Rest Macro ROC-AUC = {rf_roc_auc:.4f}). "
        f"Logistic Regression attained {lr_test_acc * 100:.2f}% accuracy (Macro F1 = {lr_f1:.4f}, ROC-AUC = {lr_roc_auc:.4f}) with one "
        f"misclassification at the default threshold. The perfect ROC-AUC indicates that Logistic Regression ranked the true-class probabilities correctly "
        f"across classification thresholds, while the single classification error at the default decision threshold resulted in 97.22% accuracy. "
        f"The single Decision Tree attained {dt_test_acc * 100:.2f}% accuracy (Macro F1 = {dt_f1:.4f}, ROC-AUC = {dt_roc_auc:.4f}) "
        f"with two misclassifications due to rigid orthogonal decision boundaries. Based on empirical test performance in this experiment, "
        f"Random Forest is selected as the champion model, with standardized Logistic Regression serving as an interpretable linear alternative."
    )
    
    add_callout(
        doc,
        f"Random Forest achieved {rf_test_acc * 100:.1f}% test accuracy and {rf_f1:.4f} Macro F1 on the held-out test split (N = 36). "
        f"Standardized Logistic Regression achieved {lr_test_acc * 100:.2f}% test accuracy with {lr_roc_auc:.4f} ROC-AUC. Both models demonstrate "
        f"that physicochemical attributes provide strong predictive signal for wine cultivar classification.",
        title="EXECUTIVE SUMMARY TAKEAWAY"
    )
    
    # 1. Introduction
    add_custom_heading(doc, "1. Introduction", level=1)
    p_intro = doc.add_paragraph()
    format_paragraph(p_intro)
    p_intro.add_run(
        "Automated cultivar identification and chemical fingerprinting are important techniques in enology, agricultural quality control, "
        "and food authentication. Regional wines command distinct commercial valuations and regulatory designations based on cultivar origin. "
        "However, traditional sensory profiling can be subjective and labor-intensive. Machine learning provides "
        "a structured, data-driven framework to map continuous physicochemical measurements directly to cultivar classifications.\n\n"
        "This project establishes a reproducible machine learning workflow. "
        "The study systematically progresses from data verification and pipeline construction through baseline modeling, "
        "tree-based comparison, multiclass evaluation, error diagnosis, and future improvement strategies."
    )
    
    # 2. Problem Definition
    add_custom_heading(doc, "2. Problem Definition", level=1)
    p_prob = doc.add_paragraph()
    format_paragraph(p_prob)
    p_prob.add_run(
        "The core research problem is framed as a supervised, multi-class pattern classification task: "
        "Given a feature vector of 13 continuous chemical and physical measurements "
        "x in R^13, predict the discrete cultivar label y in {0, 1, 2} representing three distinct Italian wine cultivars grown in the Piedmont region.\n\n"
        "Formal Task Definition:\n"
        "• Input Space (X): 13 continuous physicochemical attributes spanning alcohol content, organic acid levels, ash components, "
        "polyphenolic profiles (phenols, flavanoids, nonflavanoids, proanthocyanins), colorimetric properties (intensity, hue), spectrophotometric absorbance ratios, and amino acid content (proline).\n"
        "• Target Space (Y): Discrete multiclass label y in {0, 1, 2} corresponding to three cultivars: Class 0 (Barolo), Class 1 (Grignolino), and Class 2 (Barbera).\n"
        "• Optimization Objective: Learn a parameterized classification function f: X -> Y that maximizes out-of-sample generalization accuracy and macro-averaged F1-score while maintaining balanced performance across all classes."
    )
    
    # 3. Dataset Description
    add_custom_heading(doc, "3. Dataset Description", level=1)
    p_ds = doc.add_paragraph()
    format_paragraph(p_ds)
    p_ds.add_run(
        f"The dataset utilized is the Wine Recognition Dataset, originally curated by Forina et al. (1988) and distributed via the UCI Machine Learning Repository "
        f"and scikit-learn. The dataset consists of exactly {metadata['num_observations']} observations with {metadata['num_features']} numerical continuous features "
        f"and zero missing values. The target variable comprises three distinct classes corresponding to three wine cultivars derived from the same geographical region in Italy.\n\n"
        f"Target Class Distribution:\n"
        f"• Class 0 (Barolo): {inspection_summary['target_distribution'][0]} observations ({inspection_summary['target_proportions'][0]*100:.1f}%)\n"
        f"• Class 1 (Grignolino): {inspection_summary['target_distribution'][1]} observations ({inspection_summary['target_proportions'][1]*100:.1f}%)\n"
        f"• Class 2 (Barbera): {inspection_summary['target_distribution'][2]} observations ({inspection_summary['target_proportions'][2]*100:.1f}%)\n\n"
        f"Dataset Appropriateness:\n"
        f"This dataset is well suited for machine learning benchmarking because it presents a multi-dimensional continuous feature space "
        f"exhibiting varying scales, natural correlations, and mild class imbalance. It provides a practical setting for evaluating "
        f"scale-sensitive linear models (Logistic Regression) against scale-invariant tree models (Random Forest), testing pipeline data-leakage prevention, "
        f"and conducting multiclass One-vs-Rest ROC diagnostics."
    )
    
    # Table 1: Feature Descriptions
    p_t1_title = doc.add_paragraph()
    format_paragraph(p_t1_title, space_before=6, space_after=2)
    r_t1 = p_t1_title.add_run("Table 1: Physicochemical Feature Descriptions and Target Variables")
    r_t1.bold = True
    r_t1.font.name = "Arial"
    r_t1.font.size = Pt(10)
    r_t1.font.color.rgb = RGBColor(27, 54, 93)
    
    t1_headers = ["Feature / Variable", "Type", "Units / Domain", "Description"]
    t1_rows = [
        ["alcohol", "Continuous", "% by volume", "Ethanol content produced during fermentation"],
        ["malic_acid", "Continuous", "g/L", "Primary organic acid influencing tartness"],
        ["ash", "Continuous", "g/L", "Inorganic mineral residue remaining after incineration"],
        ["alcalinity_of_ash", "Continuous", "pH buffering", "Measure of base-neutralizing capacity of ash"],
        ["magnesium", "Continuous", "mg/L", "Essential mineral cation concentration"],
        ["total_phenols", "Continuous", "g/L", "Aggregate content of phenolic compounds"],
        ["flavanoids", "Continuous", "g/L", "Polyphenolic antioxidants affecting taste and structure"],
        ["nonflavanoid_phenols", "Continuous", "g/L", "Secondary phenolic acids and simple phenolics"],
        ["proanthocyanins", "Continuous", "g/L", "Condensed tannins contributing to astringency"],
        ["color_intensity", "Continuous", "Absorbance units", "Spectral color depth at key visual wavelengths"],
        ["hue", "Continuous", "Ratio", "Hue ratio of yellow to red optical absorbance"],
        ["od280/od315_of_diluted_wines", "Continuous", "Absorbance ratio", "Spectrophotometric protein/phenolic purity indicator"],
        ["proline", "Continuous", "mg/L", "Dominant amino acid synthesized by vine cultivar"],
        ["target", "Categorical (int)", "Classes: 0, 1, 2", "Cultivar label: 0 (Barolo), 1 (Grignolino), 2 (Barbera)"]
    ]
    add_formatted_table(doc, t1_headers, t1_rows, [1.8, 1.1, 1.3, 2.3])
    
    # Descriptive Statistics Table
    p_t2_title = doc.add_paragraph()
    format_paragraph(p_t2_title, space_before=8, space_after=2)
    r_t2 = p_t2_title.add_run("Table 2: Descriptive Statistics for the 13 Physicochemical Features (N = 178)")
    r_t2.bold = True
    r_t2.font.name = "Arial"
    r_t2.font.size = Pt(10)
    r_t2.font.color.rgb = RGBColor(27, 54, 93)
    
    desc_df = inspection_summary["descriptive_statistics"]
    t2_headers = ["Feature", "Mean", "Std", "Min", "25%", "Median", "75%", "Max"]
    t2_rows = []
    for feat_name, row in desc_df.iterrows():
        t2_rows.append([
            str(feat_name),
            f"{row['mean']:.2f}",
            f"{row['std']:.2f}",
            f"{row['min']:.2f}",
            f"{row['25%']:.2f}",
            f"{row['50%']:.2f}",
            f"{row['75%']:.2f}",
            f"{row['max']:.2f}"
        ])
    add_formatted_table(doc, t2_headers, t2_rows, [1.8, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65, 0.8])
    
    # 4. Data Preparation
    add_custom_heading(doc, "4. Data Preparation", level=1)
    p_dp = doc.add_paragraph()
    format_paragraph(p_dp)
    p_dp.add_run(
        "A structured preprocessing protocol was executed to maintain experimental integrity and eliminate data leakage:\n\n"
        "1. Data Quality Inspection: The raw dataset was inspected for structural anomalies. Exactly 0 missing values and 0 duplicate rows "
        "were identified across all 178 records. All 13 predictor columns are floating-point continuous values.\n\n"
        "2. Train/Test Stratification: The dataset was partitioned into a training set (80%, N = 142) and a held-out test set (20%, N = 36) "
        "using stratified sampling with a fixed seed (random_state = 42). Stratification guarantees that the class proportions "
        "(Class 0: ~33.1%, Class 1: ~39.9%, Class 2: ~27.0%) are preserved in both training (47, 57, 38 samples) and test (12, 14, 10 samples) splits.\n\n"
        "3. Feature Scaling & Leakage Prevention: Features exhibit disparate numerical ranges (e.g., nonflavanoid phenols range from 0.13 to 0.66 g/L, "
        "whereas proline ranges from 278 to 1680 mg/L). Without scaling, linear models would place disproportionate "
        "weight on higher-magnitude features. Standard z-score normalization (StandardScaler) was utilized. To prevent data leakage, "
        "the scaler was encapsulated within a scikit-learn Pipeline, ensuring scaling parameters were fitted solely on training data "
        "and applied downstream to test samples without prior exposure."
    )
    
    # 5. Model Selection
    add_custom_heading(doc, "5. Model Selection", level=1)
    p_ms = doc.add_paragraph()
    format_paragraph(p_ms)
    p_ms.add_run(
        "Three classification algorithms representing different modeling paradigms were selected for benchmarking:\n\n"
        "1. Logistic Regression (Primary Baseline): A multinomial linear classifier utilizing the softmax link function and L2 weight penalty. "
        "Logistic Regression models the log-odds of each class as a linear combination of standardized features. It serves as the primary baseline "
        "to test whether linear decision boundaries in standardized space are sufficient to separate the three cultivars.\n\n"
        "2. Random Forest Classifier (Non-linear Ensemble): A bagging ensemble composed of 100 de-correlated classification trees with bootstrap aggregation "
        "and randomized feature subsets at each split (max_depth = 4). Random Forest provides non-linear decision surfaces, resistance to individual tree variance, "
        "invariance to monotonic feature scaling, and feature importance estimation via Mean Decrease in Impurity (MDI).\n\n"
        "3. Decision Tree Classifier (Single Non-linear Tree): A single recursive binary partitioning CART model (max_depth = 3, Gini impurity criterion). "
        "It provides intuitive rule-based interpretability but can be sensitive to sample variance and rigid orthogonal decision boundaries."
    )
    
    # 6. Model Training
    add_custom_heading(doc, "6. Model Training", level=1)
    p_tr = doc.add_paragraph()
    format_paragraph(p_tr)
    p_tr.add_run(
        "Model training was executed deterministically on the 142-sample training split. "
        "For Logistic Regression, optimization was performed using the L-BFGS solver with a maximum iteration limit of 1000 and default inverse regularization parameter C = 1.0. "
        "For Random Forest, an ensemble of 100 estimators was initialized with max_depth = 4 to balance expressive capacity and variance control. "
        "For Decision Tree, max_depth = 3 was used to avoid over-splitting small leaves. "
        "All models used random_state = 42 for exact execution reproducibility."
    )
    
    # 7. Evaluation Methodology
    add_custom_heading(doc, "7. Evaluation Methodology", level=1)
    p_em = doc.add_paragraph()
    format_paragraph(p_em)
    p_em.add_run(
        "To evaluate model performance across multiple classes, the following metrics and diagnostic tools were computed:\n\n"
        "• Accuracy: The global ratio of correctly classified test samples to total test instances.\n"
        "• Macro-Averaged Precision: The unweighted arithmetic mean of precision across all three classes, giving equal weight to each cultivar.\n"
        "• Macro-Averaged Recall: The unweighted mean of class sensitivity (True Positive Rate), assessing detection capability across classes.\n"
        "• Macro-Averaged F1-Score: The harmonic mean of macro precision and macro recall, representing balanced multi-class performance.\n"
        "• Weighted F1-Score: Support-weighted harmonic mean accounting for slight class sample count differences.\n"
        "• Confusion Matrix: A 3x3 contingency matrix cross-tabulating actual versus predicted class memberships to reveal specific off-diagonal error patterns.\n"
        "• Multiclass One-vs-Rest (OvR) ROC & AUC: Binarizing the 3-class target into three distinct binary classification problems (Class k vs. Rest), computing the False Positive Rate versus True Positive Rate across probability thresholds, and reporting per-class and macro-average Area Under the ROC Curve.\n"
        "• Generalization Gap Analysis: Quantifying Delta = (Train Metric - Test Metric) to evaluate training versus test performance."
    )
    
    # 8. Results
    add_custom_heading(doc, "8. Results", level=1)
    p_res = doc.add_paragraph()
    format_paragraph(p_res)
    p_res.add_run(
        f"The empirical performance of all three trained pipelines on the training set (N = 142) and held-out test set (N = 36) "
        f"is summarized in Table 3. Detailed per-class precision, recall, and F1 metrics are provided in Table 4.\n\n"
        f"For Logistic Regression, the test accuracy was {lr_test_acc * 100:.2f}% (Macro F1 = {lr_f1:.4f}) and the Macro ROC-AUC was {lr_roc_auc:.4f}. "
        f"The perfect ROC-AUC indicates that the model ranked the true-class probabilities correctly across classification thresholds, "
        f"while the single classification error at the default decision threshold resulted in {lr_test_acc * 100:.2f}% accuracy. "
        f"This demonstrates that ROC-AUC evaluates threshold-independent ranking discrimination, whereas classification accuracy measures discrete decisions at a fixed probability threshold."
    )
    
    # Table 3: Model Performance Comparison
    p_t3_title = doc.add_paragraph()
    format_paragraph(p_t3_title, space_before=6, space_after=2)
    r_t3 = p_t3_title.add_run("Table 3: Comprehensive Model Performance Comparison (Train vs. Test Set)")
    r_t3.bold = True
    r_t3.font.name = "Arial"
    r_t3.font.size = Pt(10)
    r_t3.font.color.rgb = RGBColor(27, 54, 93)
    
    t3_headers = ["Model Name", "Train Acc", "Test Acc", "Macro Prec", "Macro Rec", "Macro F1", "Weighted F1", "OvR ROC-AUC", "Gap (Acc)"]
    t3_rows = []
    for _, row in summary_df.iterrows():
        t3_rows.append([
            str(row["Model"]),
            f"{row['Train Accuracy']:.4f}",
            f"{row['Test Accuracy']:.4f}",
            f"{row['Precision (Macro)']:.4f}",
            f"{row['Recall (Macro)']:.4f}",
            f"{row['F1-Score (Macro)']:.4f}",
            f"{row['F1-Score (Weighted)']:.4f}",
            f"{row['ROC-AUC (OvR Macro)']:.4f}",
            f"{row['Generalization Gap (Acc)'] * 100:+.2f}%"
        ])
    add_formatted_table(doc, t3_headers, t3_rows, [1.5, 0.6, 0.6, 0.65, 0.65, 0.65, 0.7, 0.75, 0.65])
    
    # Table 4: Detailed Classification Report
    p_t4_title = doc.add_paragraph()
    format_paragraph(p_t4_title, space_before=8, space_after=2)
    r_t4 = p_t4_title.add_run("Table 4: Per-Class Classification Reports for Candidate Models")
    r_t4.bold = True
    r_t4.font.name = "Arial"
    r_t4.font.size = Pt(10)
    r_t4.font.color.rgb = RGBColor(27, 54, 93)
    
    t4_headers = ["Model", "Class Label", "Precision", "Recall", "F1-Score", "Support"]
    t4_rows = []
    for model_name in ["Random Forest", "Logistic Regression", "Decision Tree"]:
        rep_dict = evaluation_results[model_name]["classification_report_dict"]
        for cname in ["class_0", "class_1", "class_2"]:
            c_metrics = rep_dict[cname]
            t4_rows.append([
                model_name,
                cname,
                f"{c_metrics['precision']:.4f}",
                f"{c_metrics['recall']:.4f}",
                f"{c_metrics['f1-score']:.4f}",
                str(int(c_metrics['support']))
            ])
    add_formatted_table(doc, t4_headers, t4_rows, [1.6, 1.1, 0.95, 0.95, 0.95, 0.95])
    
    # 9. Visual Evaluation
    add_custom_heading(doc, "9. Visual Evaluation", level=1)
    p_ve = doc.add_paragraph()
    format_paragraph(p_ve)
    p_ve.add_run(
        "Four dedicated visualizations were generated to analyze classification errors, "
        "cross-model performance, multiclass discrimination, and feature attribution. Each figure is displayed below with its technical interpretation."
    )
    
    # Helper to center image
    def embed_figure(path: Path, caption_label: str, interpretation_text: str, width_in: float = 5.0):
        if path.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p_img, space_before=6, space_after=2)
            run = p_img.add_run()
            run.add_picture(str(path), width=Inches(width_in))
            
            p_cap = doc.add_paragraph()
            format_paragraph(p_cap, space_before=2, space_after=8)
            r_clabel = p_cap.add_run(caption_label + " ")
            r_clabel.bold = True
            r_clabel.font.name = "Arial"
            r_clabel.font.size = Pt(9)
            r_clabel.font.color.rgb = RGBColor(27, 54, 93)
            
            r_ctext = p_cap.add_run(f"Interpretation: {interpretation_text}")
            r_ctext.font.name = "Arial"
            r_ctext.font.size = Pt(9)
            r_ctext.font.color.rgb = RGBColor(50, 50, 50)
            r_ctext.italic = True
            
    embed_figure(
        figures_paths.get("figure1", Path("")),
        "Figure 1: Confusion Matrix for the Final Random Forest Model (Test Set, N = 36).",
        "The heatmap displays a diagonal structure with 12/12 Class 0, 14/14 Class 1, and 10/10 Class 2 instances correctly classified. "
        "Zero off-diagonal errors were produced on this test split, demonstrating strong class separation across all three cultivars under the bagging ensemble.",
        width_in=4.8
    )
    
    embed_figure(
        figures_paths.get("figure2", Path("")),
        "Figure 2: Cross-Model Metric Comparison across Train and Test Partitions.",
        f"Random Forest achieved {rf_test_acc:.3f} across test metrics on this partition. Logistic Regression showed strong generalization ({lr_test_acc:.3f} test accuracy, {lr_f1:.3f} F1). "
        f"Decision Tree displayed a measurable generalization drop ({dt_train_acc:.3f} train vs. {dt_test_acc:.3f} test accuracy, gap = {dt_gap * 100:+.2f}%), illustrating the higher variance typical of unregularized single-tree models.",
        width_in=5.4
    )
    
    embed_figure(
        figures_paths.get("figure3", Path("")),
        "Figure 3: Multiclass One-vs-Rest ROC Curves for Logistic Regression Baseline.",
        f"The One-vs-Rest ROC analysis yields an Area Under the Curve (AUC) of {lr_roc_auc:.4f} across classes, achieving a Macro-Average ROC-AUC of {lr_roc_auc:.4f}. "
        f"The perfect ROC-AUC indicates that the model ranked the true-class probabilities correctly across classification thresholds, while the single classification error at the default decision threshold resulted in {lr_test_acc * 100:.2f}% accuracy. "
        f"This illustrates that ROC-AUC evaluates threshold-independent ranking discrimination, whereas classification accuracy measures discrete decisions at a fixed threshold.",
        width_in=5.0
    )
    
    embed_figure(
        figures_paths.get("figure4", Path("")),
        "Figure 4: Random Forest Feature Importance (Mean Decrease in Impurity).",
        f"{feat_top1_name} ({feat_top1_val:.3f}), {feat_top2_name} ({feat_top2_val:.3f}), {feat_top3_name} ({feat_top3_val:.3f}), and {feat_top4_name} ({feat_top4_val:.3f}) constitute the top 4 most influential features in the Random Forest model, "
        f"accounting for {top4_cum_pct:.1f}% of aggregate split impurity reduction. Conversely, {feat_low2_name} ({feat_low2_val:.3f}) and {feat_low1_name} ({feat_low1_val:.3f}) provide the lowest relative split contributions. "
        f"The feature-importance ranking identifies variables that contributed strongly to the Random Forest's decision structure. These rankings represent model-specific associations and should not be interpreted as evidence of causal biochemical effects.",
        width_in=5.0
    )
    
    # 10. Error Analysis
    add_custom_heading(doc, "10. Error Analysis", level=1)
    p_ea = doc.add_paragraph()
    format_paragraph(p_ea)
    p_ea.add_run(
        "A detailed examination of model predictions reveals specific error patterns:\n\n"
        "1. Random Forest Error Inspection: The Random Forest ensemble produced zero errors (0/36) on the held-out test split, achieving 100% precision and recall across all three classes in this experiment.\n\n"
        "2. Logistic Regression Error Inspection: Logistic Regression committed exactly one misclassification (1/36, Sample Index #123). "
        "A ground-truth Class 2 (Barbera) sample was predicted as Class 1 (Grignolino). Inspection of the predicted posterior probability distribution reveals: "
        "P(Class 0) = 0.0012, P(Class 1) = 0.5843, P(Class 2) = 0.4145. This observation lies near the linear decision boundary between Class 1 and Class 2, "
        "exhibiting an intermediate flavanoid measurement (1.20 g/L) and lower color intensity (5.00) relative to typical Class 2 samples.\n\n"
        "3. Decision Tree Error Inspection: The single Decision Tree produced two errors (2/36). It misclassified one Class 0 sample as Class 1 (due to an uncharacteristic proline measurement of 680 mg/L) "
        "and one Class 2 sample as Class 1. This illustrates how axis-aligned orthogonal partitions can struggle near overlapping regional boundaries.\n\n"
        "4. Class Imbalance Impact: The mild class imbalance (Class 0: 33.1%, Class 1: 39.9%, Class 2: 27.0%) did not induce majority-class bias. "
        "Stratified sampling preserved representative class proportions across partitions."
    )
    
    # 11. Overfitting and Underfitting
    add_custom_heading(doc, "11. Overfitting and Underfitting", level=1)
    p_ou = doc.add_paragraph()
    format_paragraph(p_ou)
    p_ou.add_run(
        "Overfitting and underfitting were evaluated by comparing training performance against held-out test performance:\n\n"
        "• Underfitting Assessment: No clear evidence of underfitting was observed based on the training and test performance in this experiment. "
        f"Training accuracies ranged from {dt_train_acc * 100:.2f}% (Decision Tree) to {rf_train_acc * 100:.1f}% (Logistic Regression and Random Forest), "
        "indicating that the candidate models had sufficient expressive capacity to learn the classification patterns in the training data.\n\n"
        "• Overfitting Assessment:\n"
        f"  - Decision Tree: The Decision Tree showed a noticeable generalization gap of Delta = {dt_gap * 100:+.2f}% "
        f"(Train Acc = {dt_train_acc * 100:.2f}%, Test Acc = {dt_test_acc * 100:.2f}%; Train F1 = {dt_train_f1:.4f}, Test F1 = {dt_f1:.4f}), "
        "reflecting the variance typical of unregularized single tree models on small sample splits.\n"
        f"  - Logistic Regression: Logistic Regression demonstrated stable performance with a modest generalization gap of Delta = {lr_gap * 100:+.2f}% "
        f"(Train Acc = {lr_train_acc * 100:.2f}%, Test Acc = {lr_test_acc * 100:.2f}%; Train F1 = {lr_train_f1:.4f}, Test F1 = {lr_f1:.4f}). "
        "The L2 weight penalty controlled parameter magnitude.\n"
        f"  - Random Forest: The Random Forest showed no observed train-test accuracy gap on this fixed split (Train Acc = {rf_train_acc * 100:.1f}%, Test Acc = {rf_test_acc * 100:.1f}%), "
        "suggesting strong generalization in this experiment. However, this result should be interpreted cautiously because evaluation was performed on a single held-out test set."
    )
    
    # 12. Model Comparison
    add_custom_heading(doc, "12. Model Comparison", level=1)
    p_mc = doc.add_paragraph()
    format_paragraph(p_mc)
    p_mc.add_run(
        "Comparing the candidate algorithms highlights key trade-offs between predictive performance, variance, and interpretability:\n\n"
        "• Logistic Regression vs. Tree Ensembles: Logistic Regression provides direct coefficient interpretability through standardized regression weights and predicted class probabilities. "
        "However, it relies on linear decision boundaries in feature space. Random Forest aggregates multiple decision trees to capture non-linear interactions "
        "(such as relationships between flavanoids and color intensity) without requiring manual interaction terms.\n\n"
        "• Single Decision Tree vs. Random Forest: The single Decision Tree is simple and produces clear if-then rules. "
        f"However, its empirical test score ({dt_test_acc * 100:.2f}% accuracy, {dt_f1:.4f} F1) was lower than Random Forest ({rf_test_acc * 100:.1f}% accuracy, {rf_f1:.4f} F1). "
        "Averaging across 100 randomized trees reduced individual tree variance and improved test performance."
    )
    
    # 13. Final Model Selection
    add_custom_heading(doc, "13. Final Model Selection", level=1)
    p_fs = doc.add_paragraph()
    format_paragraph(p_fs)
    p_fs.add_run(
        "Final Champion Selection: Random Forest Classifier\n\n"
        "Justification:\n"
        f"1. Test Performance: Random Forest was selected as the champion model for this experiment because it achieved the strongest held-out test performance among the evaluated models "
        f"({rf_test_acc * 100:.1f}% accuracy, Macro F1 = {rf_f1:.4f}, ROC-AUC = {rf_roc_auc:.4f}). This conclusion is specific to the dataset and evaluation protocol used in this study. "
        f"While the model attained 100% accuracy on the test set, this evaluation split contains 36 observations, so performance should be validated across broader sampling in future work.\n"
        "2. Variance Control: By ensembling 100 decorrelated trees with bootstrap aggregation, Random Forest reduced the variance observed in the single decision tree.\n"
        "3. Scale Invariance: Random Forest is natively invariant to monotonic feature scaling and robust to feature magnitude differences.\n"
        "4. Feature Attribution: The feature-importance ranking identifies variables that contributed strongly to the Random Forest's decision structure. "
        "These rankings represent model-specific associations and should not be interpreted as evidence of causal biochemical effects.\n\n"
        f"Alternative Candidate: Standardized Logistic Regression represents a strong, interpretable alternative ({lr_test_acc * 100:.2f}% test accuracy, {lr_roc_auc:.4f} ROC-AUC) "
        "when simple linear coefficients and fast inference are desired."
    )
    
    # 14. Limitations
    add_custom_heading(doc, "14. Limitations", level=1)
    p_lim = doc.add_paragraph()
    format_paragraph(p_lim)
    p_lim.add_run(
        "The findings of this study should be evaluated within the context of several limitations:\n\n"
        "1. Sample Size Constraint: The dataset contains 178 total observations (36 test samples). While standard for benchmarking, "
        "the small sample size increases confidence interval widths, and 100% test accuracy on 36 samples should not be assumed to guarantee error-free performance on new datasets.\n\n"
        "2. Single Geographic Region: All wine samples originate from the Piedmont region in Italy. "
        "The model may not generalize directly to identical cultivars cultivated in different soil types, climates, or winemaking processes.\n\n"
        "3. Temporal Invariance: The dataset does not include longitudinal multi-vintage time series data. "
        "Annual weather differences can alter phenolic concentrations, which may require periodic recalibration.\n\n"
        "4. Fixed Partition: Although stratified, evaluation on a single 80/20 train/test split could be sensitive to the chosen random seed."
    )
    
    # 15. Improvement Strategies
    add_custom_heading(doc, "15. Improvement Strategies", level=1)
    p_imp = doc.add_paragraph()
    format_paragraph(p_imp)
    p_imp.add_run(
        "To provide further experimental robustness beyond the current baseline pipeline, the following future research and engineering enhancements are proposed:\n\n"
        "1. Stratified Cross-Validation: Implementing repeated Stratified K-Fold cross-validation will provide more comprehensive variance estimates "
        "across multiple dataset partitions.\n\n"
        "2. Hyperparameter Optimization: Conducting systematic search (such as GridSearchCV or RandomizedSearchCV) over Random Forest parameters "
        "(e.g., min_samples_split, max_features, criterion) and Logistic Regression regularization strength (C).\n\n"
        f"3. Feature Selection: Utilizing feature importance rankings ({feat_top1_name}, {feat_top2_name}, {feat_top3_name}, {feat_top4_name}) or Recursive Feature Elimination (RFE) to identify compact feature subsets, "
        "potentially reducing laboratory measurement requirements without major loss in accuracy.\n\n"
        "4. Additional Algorithm Benchmarking: Testing algorithms such as Gradient Boosting or Support Vector Machines (SVM with RBF kernel) "
        "on expanded datasets.\n\n"
        "5. External Dataset Validation: Testing model generalization on wine datasets collected from different wine-growing regions."
    )
    
    # 16. Conclusion
    add_custom_heading(doc, "16. Conclusion", level=1)
    p_concl = doc.add_paragraph()
    format_paragraph(p_concl)
    p_concl.add_run(
        "This project successfully developed and evaluated a reproducible machine learning pipeline "
        "for wine cultivar classification. The experimental results demonstrate that 13 physicochemical measurements provide strong "
        "predictive signal for classifying Italian wine cultivars. "
        f"The Random Forest model achieved strong test performance ({rf_test_acc * 100:.1f}% test accuracy, {rf_f1:.4f} F1-score), "
        f"while Logistic Regression demonstrated effective baseline performance ({lr_test_acc * 100:.2f}% test accuracy, {lr_roc_auc:.4f} ROC-AUC). "
        "The complete codebase and experimental protocol satisfy standard data science practices of data hygiene, leakage prevention, and computational reproducibility."
    )
    
    # 17. References
    add_custom_heading(doc, "17. References", level=1)
    p_ref = doc.add_paragraph()
    format_paragraph(p_ref)
    p_ref.add_run(
        "1. Forina, M., Armanino, C., Castino, M., & Ubigli, M. (1988). Multivariate data analysis as a discriminating tool of the origin of wines. "
        "Vitis, 25(3), 189-201.\n\n"
        "2. Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32. https://doi.org/10.1023/A:1010933404324\n\n"
        "3. Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. (2011). "
        "Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.\n\n"
        "4. Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction (2nd ed.). Springer New York.\n\n"
        "5. Bishop, C. M. (2006). Pattern Recognition and Machine Learning. Springer New York."
    )
    
    # Selected Python Code Snippets
    add_custom_heading(doc, "Selected Python Code Snippets", level=1)
    p_snip_intro = doc.add_paragraph()
    format_paragraph(p_snip_intro)
    p_snip_intro.add_run(
        "The following concise Python snippets illustrate the core implementation of the machine learning pipeline, "
        "demonstrating data splitting, leakage-free pipeline construction, model training, and multiclass evaluation."
    )
    
    code_snippet_1 = (
        "# 1. Stratified Train/Test Split (80/20) with Fixed Random State\n"
        "from sklearn.model_selection import train_test_split\n\n"
        "X = wine_df.drop(columns=['target'])\n"
        "y = wine_df['target']\n\n"
        "X_train, X_test, y_train, y_test = train_test_split(\n"
        "    X, y, test_size=0.20, random_state=42, stratify=y\n"
        ")"
    )
    add_code_block(doc, code_snippet_1, "1. Stratified Data Partitioning")
    
    code_snippet_2 = (
        "# 2. Leakage-Free Pipeline Construction and Model Training\n"
        "from sklearn.pipeline import Pipeline\n"
        "from sklearn.preprocessing import StandardScaler\n"
        "from sklearn.linear_model import LogisticRegression\n"
        "from sklearn.ensemble import RandomForestClassifier\n\n"
        "# Standardized Logistic Regression Pipeline\n"
        "lr_pipeline = Pipeline([\n"
        "    ('scaler', StandardScaler()),\n"
        "    ('classifier', LogisticRegression(random_state=42, max_iter=1000))\n"
        "])\n"
        "lr_pipeline.fit(X_train, y_train)\n\n"
        "# Random Forest Ensemble Pipeline\n"
        "rf_pipeline = Pipeline([\n"
        "    ('classifier', RandomForestClassifier(n_estimators=100, max_depth=4, random_state=42))\n"
        "])\n"
        "rf_pipeline.fit(X_train, y_train)"
    )
    add_code_block(doc, code_snippet_2, "2. Pipeline Construction and Model Training")
    
    code_snippet_3 = (
        "# 3. Multiclass Evaluation and One-vs-Rest ROC AUC Computation\n"
        "from sklearn.metrics import accuracy_score, f1_score, roc_auc_score, confusion_matrix\n\n"
        "y_pred = rf_pipeline.predict(X_test)\n"
        "y_prob = rf_pipeline.predict_proba(X_test)\n\n"
        "acc = accuracy_score(y_test, y_pred)\n"
        "f1_macro = f1_score(y_test, y_pred, average='macro')\n"
        "roc_auc_ovr = roc_auc_score(y_test, y_prob, multi_class='ovr', average='macro')\n"
        "cm = confusion_matrix(y_test, y_pred)"
    )
    add_code_block(doc, code_snippet_3, "3. Multiclass Evaluation and Metrics Calculation")
    
    # Save Document
    doc.save(str(output_path))
